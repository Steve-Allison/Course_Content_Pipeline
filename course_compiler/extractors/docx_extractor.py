"""
Microsoft Word (.docx) file extractor implementation.
"""
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from docx import Document as DocxDocument
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import io
import os

from .base import BaseExtractor, ExtractionError
from ..logging_utils import get_logger

logger = get_logger(__name__)

class DocxExtractor(BaseExtractor):
    """Extracts content from Microsoft Word (.docx) files."""

    def __init__(self, file_path: Path, config: Optional[Dict[str, Any]] = None):
        super().__init__(file_path, config)
        self.document: Optional[DocumentType] = None
        self.images_dir = self.config.get('images_dir', 'images')

    def load(self) -> None:
        """Load the Word document."""
        try:
            self.document = DocxDocument(self.file_path)
        except Exception as e:
            raise ExtractionError(f"Failed to load Word document: {e}") from e

    def extract(self) -> Dict[str, Any]:
        """
        Extract content from the Word document.

        Returns:
            Dictionary containing paragraphs, tables, and metadata
        """
        if self.document is None:
            self.load()

        try:
            metadata = self.get_metadata()

            # Create output directory for images if it doesn't exist
            os.makedirs(self.images_dir, exist_ok=True)

            # Process document content
            content = {
                'metadata': metadata,
                'paragraphs': [],
                'tables': [],
                'images': [],
                'sections': [],
                'file_type': 'docx'
            }

            # Extract sections and their content
            for section in self.document.sections:
                section_data = {
                    'start_page': len(content['paragraphs']) + 1,  # Approximate
                    'header': self._extract_header_footer(section.header),
                    'footer': self._extract_header_footer(section.footer),
                    'page_width': section.page_width,
                    'page_height': section.page_height,
                    'orientation': 'landscape' if section.orientation == 1 else 'portrait'
                }
                content['sections'].append(section_data)

            # Process all elements in the document
            for element in self.document.element.body.iterchildren():
                if isinstance(element, CT_P):  # Paragraph
                    paragraph = self._process_paragraph(element)
                    if paragraph:
                        content['paragraphs'].append(paragraph)
                elif isinstance(element, CT_Tbl):  # Table
                    table = self._process_table(element)
                    if table:
                        content['tables'].append(table)

            # Extract images
            content['images'] = self._extract_images()

            return content

        except Exception as e:
            logger.error(f"Error extracting content from {self.file_path}: {e}")
            raise ExtractionError(f"Failed to extract content: {e}") from e

    def _process_paragraph(self, element: CT_P) -> Optional[Dict[str, Any]]:
        """Process a single paragraph element."""
        try:
            paragraph = self.document.paragraphs[
                self.document.paragraphs.index(element.paragraph)
            ]

            if not paragraph.text.strip():
                return None

            return {
                'text': paragraph.text.strip(),
                'style': paragraph.style.name if paragraph.style else 'Normal',
                'runs': [
                    {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None
                    }
                    for run in paragraph.runs
                ]
            }
        except Exception as e:
            logger.warning(f"Failed to process paragraph: {e}")
            return None

    def _process_table(self, element: CT_Tbl) -> Optional[Dict[str, Any]]:
        """Process a table element."""
        try:
            table = self.document.tables[
                [t._tbl for t in self.document.tables].index(element)
            ]

            rows = []
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = ' '.join(para.text for para in cell.paragraphs)
                    row_data.append(cell_text.strip())
                rows.append(row_data)

            return {
                'rows': rows,
                'dimensions': (len(rows), len(rows[0])) if rows else (0, 0)
            }
        except Exception as e:
            logger.warning(f"Failed to process table: {e}")
            return None

    def _extract_header_footer(self, section_part) -> Dict[str, Any]:
        """Extract content from header or footer."""
        if not section_part:
            return {}

        content = {
            'paragraphs': [],
            'tables': []
        }

        for element in section_part.paragraphs:
            if element.text.strip():
                content['paragraphs'].append(element.text.strip())

        for table in section_part.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)

        return content

    def _extract_images(self) -> List[Dict[str, Any]]:
        """Extract images from the document."""
        images = []
        try:
            # This is a simplified version - actual image extraction requires more complex handling
            # as docx stores images in a binary format within the document
            for rel in self.document.part.rels.values():
                if 'image' in rel.reltype:
                    img_data = rel.target_part.blob
                    if img_data:
                        image_name = f"docx_img_{hash(img_data) & 0xFFFFFFFF}.png"
                        image_path = os.path.join(self.images_dir, image_name)

                        with open(image_path, 'wb') as f:
                            f.write(img_data)

                        images.append({
                            'path': image_path,
                            'name': image_name,
                            'size': len(img_data)
                        })
        except Exception as e:
            logger.warning(f"Failed to extract images: {e}")

        return images
